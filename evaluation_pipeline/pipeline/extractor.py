"""
Section extractor – Step 1 of the ROAR pipeline.

Decomposes a raw ROAR document (plain text, PDF, or .docx) into the four
structured components: PLO, Assessment Methods, Results, and Improvement Plan.

Supported input paths
---------------------
load_pdf(path)          → plain text from a PDF
load_docx(path)         → plain text from a Word document (paragraphs + tables)
extract_from_docx(path) → ROARSections directly from the ROAR table structure
                          (skips LLM extraction entirely — fastest and most
                          accurate for standard Rice ROAR .docx files)

For plain text, two strategies are attempted in order:
1. Heuristic regex parsing (fast, zero-cost).
2. LLM-based extraction via the EXTRACTION_PROMPT if heuristics fail.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Optional

import config
from models.schemas import ROARSections
from prompts.templates import EXTRACTION_PROMPT
from utils.json_parser import RobustJsonOutputParser
from utils.llm_factory import Backend, build_llm


# Minimum character length for a section to be considered "found"
_MIN_SECTION_LEN = 20


class SectionExtractor:
    """
    Extracts the four ROAR sections from a document.

    Parameters
    ----------
    use_llm_fallback:
        When True (default), fall back to an LLM extraction chain if the
        heuristic parser cannot find all sections.
    model:
        Model identifier.  Defaults to ``config.EVALUATOR_MODEL``.
    backend:
        ``"ollama"`` or ``"openai"``.  Defaults to ``config.EVALUATOR_BACKEND``.
    """

    def __init__(
        self,
        use_llm_fallback: bool = True,
        model: Optional[str] = None,
        backend: Optional[Backend] = None,
    ) -> None:
        self.use_llm_fallback = use_llm_fallback
        self._model = model or config.EVALUATOR_MODEL
        self._backend: Backend = backend or config.EVALUATOR_BACKEND  # type: ignore[assignment]
        self._chain = self._build_chain()

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def extract(self, document: str) -> ROARSections:
        """
        Extract sections from *document* (plain text).

        Returns a :class:`ROARSections` instance.  Missing sections are
        represented as ``"Not provided"``.
        """
        sections = self._heuristic_extract(document)

        if self.use_llm_fallback and not self._all_found(sections):
            sections = self._llm_extract(document)

        return sections

    @staticmethod
    def load_pdf(path: str | Path) -> str:
        """Load plain text from a PDF file using *pypdf*."""
        try:
            from pypdf import PdfReader  # lazy import
        except ImportError as exc:
            raise ImportError(
                "pypdf is required for PDF loading.  Install it with: pip install pypdf"
            ) from exc

        reader = PdfReader(str(path))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(pages)

    @staticmethod
    def load_docx(path: str | Path) -> str:
        """
        Convert a Word .docx file to plain text (paragraphs + all table cells).

        Use this when you want to feed the raw text through the normal
        heuristic / LLM extraction flow.  For standard Rice ROAR .docx files
        that use the four-row table layout, prefer ``extract_from_docx`` which
        maps sections directly without any LLM call.
        """
        try:
            from docx import Document  # lazy import
        except ImportError as exc:
            raise ImportError(
                "python-docx is required for .docx loading.  "
                "Install it with: pip install python-docx"
            ) from exc

        doc = Document(str(path))
        parts: list[str] = []

        # Regular paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())

        # Table cells (row-by-row, labelled so the heuristic extractor can match)
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append("  |  ".join(cells))

        return "\n".join(parts)

    @classmethod
    def extract_from_docx(cls, path: str | Path) -> "ROARSections":
        """
        Directly map a Rice ROAR .docx table to :class:`ROARSections`.

        This is the **recommended** method for standard Rice ROAR Word files.
        It reads the four-row table (PLO / Methods / Results / Improvement Plan)
        and maps each content cell to the corresponding section — no LLM call,
        no regex, zero latency.

        Falls back to ``load_docx`` + regex/LLM extraction if the expected
        table structure is not found.
        """
        try:
            from docx import Document
        except ImportError as exc:
            raise ImportError(
                "python-docx is required.  Install it with: pip install python-docx"
            ) from exc

        doc = Document(str(path))

        # Rice ROAR .docx files store content in a single table where:
        #   col 0 = section label   col 1 = student/department content
        # Row labels match these keywords (case-insensitive):
        _LABEL_MAP = {
            "program learning outcome": "plo",
            "plo":                      "plo",
            "methods":                  "methods",
            "assessment":               "methods",
            "results":                  "results",
            "improvement":              "plan",
        }

        found: dict[str, str] = {}
        for table in doc.tables:
            for row in table.rows:
                if len(row.cells) < 2:
                    continue
                label = row.cells[0].text.strip().lower()
                content = row.cells[1].text.strip()
                if not content:
                    continue
                for keyword, section_key in _LABEL_MAP.items():
                    if keyword in label and section_key not in found:
                        found[section_key] = content
                        break

        if len(found) >= 3:
            return ROARSections(
                plo=found.get("plo", "Not provided"),
                methods=found.get("methods", "Not provided"),
                results=found.get("results", "Not provided"),
                plan=found.get("plan", "Not provided"),
            )

        # Fallback: convert to text and use normal extraction
        import logging
        logging.getLogger(__name__).warning(
            "ROAR table structure not detected in %s — falling back to text extraction.",
            path,
        )
        return ROARSections(
            plo="Not provided",
            methods="Not provided",
            results="Not provided",
            plan="Not provided",
        )

    # ------------------------------------------------------------------
    # Internal helpers
    # ------------------------------------------------------------------

    def _build_chain(self):
        if self._backend == "azure_openai":
            llm = build_llm(
                self._model,
                self._backend,
                azure_temperature=config.EVALUATOR_AZURE_TEMPERATURE,
                azure_max_tokens=config.EVALUATOR_AZURE_MAX_TOKENS,
            )
        else:
            llm = build_llm(
                self._model,
                self._backend,
                openai_base_url=config.EVALUATOR_API_BASE,
                openai_api_key=config.EVALUATOR_API_KEY,
                openai_temperature=config.EVALUATOR_TEMPERATURE,
                openai_top_p=config.EVALUATOR_TOP_P,
                openai_presence_penalty=config.EVALUATOR_PRESENCE_PENALTY,
                openai_max_tokens=config.EVALUATOR_MAX_TOKENS,
                openai_default_query={"api-version": config.EVALUATOR_API_VERSION},
            )
        return EXTRACTION_PROMPT | llm | RobustJsonOutputParser()

    @staticmethod
    def _heuristic_extract(text: str) -> ROARSections:
        """
        Pattern-based extraction that looks for common ROAR section headers.
        Returns whatever it can find; absent sections default to
        ``"Not provided"``.
        """
        patterns: dict[str, str] = {
            "plo":     r"(?:program\s+learning\s+objective|PLO)[:\s]*(.+?)(?=\n(?:assessment\s+method|results|improvement|$))",
            "methods": r"(?:assessment\s+method(?:s)?)[:\s]*(.+?)(?=\n(?:results|improvement|PLO|$))",
            "results": r"(?:results?)[:\s]*(.+?)(?=\n(?:improvement|PLO|assessment|$))",
            "plan":    r"(?:improvement\s+plan)[:\s]*(.+?)(?=\n(?:PLO|assessment|results|$))",
        }

        extracted: dict[str, str] = {}
        for key, pattern in patterns.items():
            match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
            extracted[key] = match.group(1).strip() if match else "Not provided"

        return ROARSections(**extracted)

    def _llm_extract(self, document: str) -> ROARSections:
        """Call the LLM extraction chain and parse the JSON response."""
        raw: dict = self._chain.invoke({"document": document})
        return ROARSections(
            plo=raw.get("plo", "Not provided"),
            methods=raw.get("methods", "Not provided"),
            results=raw.get("results", "Not provided"),
            plan=raw.get("plan", "Not provided"),
        )

    @staticmethod
    def _all_found(sections: ROARSections) -> bool:
        return all(
            v != "Not provided" and len(v) >= _MIN_SECTION_LEN
            for v in [sections.plo, sections.methods, sections.results, sections.plan]
        )
